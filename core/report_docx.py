# Word 报告生成（python-docx，零系统依赖）— D8.3
#
# 文档结构：标题 → 摘要表格 → 风险分布表格 → 去重统计 → 漏洞详情表 → 其他结果 → 修复建议汇总
# 中文字体：通过 run.font.name + eastAsia 设置，Word/LibreOffice 会自动回退替换
import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.models import (STATUS_CONFIRMED, STATUS_SAFE, STATUS_UNKNOWN,
                          SEVERITY_HIGH, SEVERITY_MEDIUM, SEVERITY_LOW,
                          SEVERITY_CN)

# 中文字体（Word/LibreOffice 会自动回退到系统可用字体）
_FONT = '微软雅黑'
_FONT_MONO = 'Consolas'

# 严重度颜色（RGB）
_SEV_COLOR = {
    SEVERITY_HIGH: RGBColor(0xcf, 0x22, 0x2e),
    SEVERITY_MEDIUM: RGBColor(0x9a, 0x67, 0x00),
    SEVERITY_LOW: RGBColor(0x1a, 0x7f, 0x37),
}
# 严重度单元格背景色（ARGB hex for shading）
_SEV_FILL = {
    SEVERITY_HIGH: 'CF222E',
    SEVERITY_MEDIUM: 'D4A72C',
    SEVERITY_LOW: '1A7F37',
}
# 状态中文名
_STATUS_CN = {
    STATUS_CONFIRMED: '已确认',
    STATUS_SAFE: '安全',
    STATUS_UNKNOWN: '未知',
}


def _set_run_font(run, name=_FONT, size=None, color=None, bold=None, mono=False):
    """设置 run 字体（中文需同时设 ascii 和 eastAsia）

    Args:
        run: docx run 对象
        name: 字体名（默认微软雅黑）
        size: 字号 Pt
        color: RGBColor
        bold: 是否加粗
        mono: True 时用等宽字体（URL/证据列）
    """
    font_name = _FONT_MONO if mono else name
    run.font.name = font_name
    # 设置中文字体（eastAsia）— python-docx 必须通过 XML 操作
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), name if not mono else _FONT_MONO)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold


def _set_cell_bg(cell, hex_color):
    """设置表格单元格背景色

    Args:
        cell: docx cell 对象
        hex_color: 6 位十六进制颜色（如 'CF222E'）
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = tc_pr.makeelement(qn('w:shd'), {})
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)


def _add_cell_text(cell, text, size=9, bold=False, color=None, mono=False, align=None):
    """向单元格添加文本（清空默认空段落后写入）

    Args:
        cell: docx cell 对象
        text: 文本内容
        size: 字号
        bold: 加粗
        color: RGBColor
        mono: 等宽字体
        align: WD_ALIGN_PARAGRAPH 常量
    """
    # 清空默认段落
    cell.text = ''
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text) if text is not None else '')
    _set_run_font(run, size=size, bold=bold, color=color, mono=mono)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _add_heading(doc, text, level=1, size=None, color=None):
    """添加标题段落（自定义样式，避免依赖默认 Heading 样式的中文字体问题）"""
    para = doc.add_paragraph()
    if level == 0:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    default_size = {0: 22, 1: 16, 2: 13, 3: 11}.get(level, 11)
    default_color = RGBColor(0x1f, 0x23, 0x28)
    _set_run_font(run, size=size or default_size, bold=True,
                  color=color or default_color)
    para.paragraph_format.space_before = Pt(12 if level > 0 else 0)
    para.paragraph_format.space_after = Pt(6)
    return para


def _style_table_header(table, fill='0969DA'):
    """设置表格首行为表头样式（蓝底白字加粗）"""
    for cell in table.rows[0].cells:
        _set_cell_bg(cell, fill)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                _set_run_font(run, size=9, bold=True,
                              color=RGBColor(0xff, 0xff, 0xff))


def _set_table_borders(table):
    """为表格添加细边框（python-docx 默认无边框）"""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = borders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = borders.makeelement(qn(f'w:{edge}'), {})
            borders.append(elem)
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')       # 0.5pt
        elem.set(qn('w:color'), 'D0D7DE')


def render_docx(builder, out_path):
    """渲染 Word 报告到 out_path

    Args:
        builder: ReportBuilder 实例（使用 _effective_results() 获取去重后结果）
        out_path: 输出文件路径
    Returns:
        out_path
    """
    doc = Document()

    # === 设置默认样式字体 ===
    style = doc.styles['Normal']
    style.font.name = _FONT
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), _FONT)
    rFonts.set(qn('w:hAnsi'), _FONT)
    rFonts.set(qn('w:eastAsia'), _FONT)

    # === 封面标题 ===
    _add_heading(doc, '若依综合漏洞检测报告', level=0, size=24)

    # 摘要信息
    target = builder.target or '未指定'
    scan_time = builder.summary.get('started_at', '')
    if isinstance(scan_time, (int, float)) and scan_time:
        scan_time = datetime.fromtimestamp(scan_time).strftime('%Y-%m-%d %H:%M:%S')
    if not scan_time:
        scan_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    gen_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    duration = builder.summary.get('duration', 0)
    req_count = builder.summary.get('request_count', 0)
    mode = builder.summary.get('mode', '')
    fp = builder.summary.get('fingerprint', {}) or {}
    fp_cms = fp.get('cms', '')
    fp_conf = fp.get('confidence', 0)

    # === 摘要表格 ===
    _add_heading(doc, '扫描摘要', level=2)
    summary_rows = [
        ('扫描目标', target),
        ('CMS 识别', f'{fp_cms or "未识别"}（置信度 {fp_conf:.2f}）' if fp_cms else '未识别'),
        ('扫描模式', str(mode) if mode else '未指定'),
        ('扫描时间', str(scan_time)),
        ('报告生成', gen_time),
        ('耗时（秒）', f'{duration:.2f}' if isinstance(duration, (int, float)) else str(duration)),
        ('请求数', str(req_count)),
    ]
    summary_table = doc.add_table(rows=len(summary_rows), cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(summary_table)
    for i, (k, v) in enumerate(summary_rows):
        _add_cell_text(summary_table.rows[i].cells[0], k, size=10, bold=True,
                       color=RGBColor(0x1f, 0x23, 0x28))
        _set_cell_bg(summary_table.rows[i].cells[0], 'F6F8FA')
        _add_cell_text(summary_table.rows[i].cells[1], v, size=10)
    # 设置列宽
    for row in summary_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    # === 风险分布 ===
    _add_heading(doc, '风险分布（仅统计确认漏洞）', level=2)
    dist = builder.risk_distribution()
    dist_table = doc.add_table(rows=2, cols=5)
    dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(dist_table)
    headers = ['高危', '中危', '低危', '合计', '']
    values = [str(dist['high']), str(dist['medium']), str(dist['low']),
              str(dist['total']), '']
    fills = [_SEV_FILL[SEVERITY_HIGH], _SEV_FILL[SEVERITY_MEDIUM],
             _SEV_FILL[SEVERITY_LOW], '656D76', '']
    for j, (h, v, f) in enumerate(zip(headers, values, fills)):
        _add_cell_text(dist_table.rows[0].cells[j], h, size=10, bold=True,
                       color=RGBColor(0xff, 0xff, 0xff),
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        if f:
            _set_cell_bg(dist_table.rows[0].cells[j], f)
        _add_cell_text(dist_table.rows[1].cells[j], v, size=12, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

    # === 去重统计 ===
    dr = builder.dedup_report()
    if dr and dr.merged_groups > 0:
        dedup_para = doc.add_paragraph()
        dedup_run = dedup_para.add_run(
            f'去重统计：原始 {dr.original_count} 条 → 聚合后 {dr.aggregated_count} 条'
            f'（{dr.merged_groups} 组合并）'
        )
        _set_run_font(dedup_run, size=10, color=RGBColor(0x65, 0x6d, 0x76))
        dedup_para.paragraph_format.space_before = Pt(6)

    # === 漏洞详情表 ===
    _add_heading(doc, '漏洞详情', level=1)
    confirmed = builder.confirmed_results()
    if not confirmed:
        empty_para = doc.add_paragraph()
        empty_run = empty_para.add_run('未发现确认漏洞。')
        _set_run_font(empty_run, size=10, color=RGBColor(0x65, 0x6d, 0x76))
    else:
        headers = ['漏洞名称', '严重度', 'URL', '证据', '修复建议']
        vuln_table = doc.add_table(rows=len(confirmed) + 1, cols=len(headers))
        vuln_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(vuln_table)
        # 表头
        for j, h in enumerate(headers):
            _add_cell_text(vuln_table.rows[0].cells[j], h, size=9, bold=True,
                           color=RGBColor(0xff, 0xff, 0xff),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(vuln_table.rows[0].cells[j], '0969DA')
        # 数据行
        for i, r in enumerate(confirmed, 1):
            sev_cn = SEVERITY_CN.get(r.severity, r.severity)
            hit_info = ''
            if hasattr(r, 'hit_count') and r.hit_count > 1:
                hit_info = f' (x{r.hit_count})'
            name_color = _SEV_COLOR.get(r.severity)
            _add_cell_text(vuln_table.rows[i].cells[0], r.name + hit_info,
                           size=9, bold=True, color=name_color)
            _add_cell_text(vuln_table.rows[i].cells[1], sev_cn, size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_cell_text(vuln_table.rows[i].cells[2], r.url or '', size=8, mono=True)
            _add_cell_text(vuln_table.rows[i].cells[3], r.evidence or '', size=8, mono=True)
            _add_cell_text(vuln_table.rows[i].cells[4], r.fix or '', size=9)
            # 隔行底色
            if i % 2 == 0:
                for cell in vuln_table.rows[i].cells:
                    _set_cell_bg(cell, 'F6F8FA')
        # 列宽
        col_widths = [Cm(3.5), Cm(1.5), Cm(4), Cm(4), Cm(3.5)]
        for row in vuln_table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

    # === 其他结果（非 CONFIRMED）===
    all_results = builder._effective_results()
    others = [r for r in all_results if r.status != STATUS_CONFIRMED]
    if others:
        doc.add_page_break()
        _add_heading(doc, '其他结果（未确认/安全）', level=1)
        other_headers = ['名称', '状态', 'URL', '证据']
        other_table = doc.add_table(rows=len(others) + 1, cols=len(other_headers))
        other_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(other_table)
        for j, h in enumerate(other_headers):
            _add_cell_text(other_table.rows[0].cells[j], h, size=9, bold=True,
                           color=RGBColor(0xff, 0xff, 0xff),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(other_table.rows[0].cells[j], '656D76')
        for i, r in enumerate(others, 1):
            _add_cell_text(other_table.rows[i].cells[0], r.name, size=9)
            _add_cell_text(other_table.rows[i].cells[1],
                           _STATUS_CN.get(r.status, r.status), size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_cell_text(other_table.rows[i].cells[2], r.url or '', size=8, mono=True)
            _add_cell_text(other_table.rows[i].cells[3], r.evidence or '', size=8, mono=True)
            if i % 2 == 0:
                for cell in other_table.rows[i].cells:
                    _set_cell_bg(cell, 'F6F8FA')
        col_widths = [Cm(4), Cm(2), Cm(5), Cm(5)]
        for row in other_table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

    # === 修复建议汇总 ===
    if confirmed:
        doc.add_page_break()
        _add_heading(doc, '修复建议汇总', level=1)
        for i, r in enumerate(confirmed, 1):
            sev_cn = SEVERITY_CN.get(r.severity, r.severity)
            title_color = _SEV_COLOR.get(r.severity, RGBColor(0x1f, 0x23, 0x28))
            # 标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f'{i}. {r.name}（{sev_cn}）')
            _set_run_font(title_run, size=11, bold=True, color=title_color)
            title_para.paragraph_format.space_before = Pt(10)
            # URL
            url_para = doc.add_paragraph()
            url_label = url_para.add_run('URL：')
            _set_run_font(url_label, size=10, bold=True)
            url_val = url_para.add_run(r.url or '无')
            _set_run_font(url_val, size=10, mono=True)
            # 修复建议
            fix_para = doc.add_paragraph()
            fix_label = fix_para.add_run('修复：')
            _set_run_font(fix_label, size=10, bold=True,
                          color=RGBColor(0x1a, 0x7f, 0x37))
            fix_val = fix_para.add_run(r.fix or '暂无建议')
            _set_run_font(fix_val, size=10)

    # === 页脚 ===
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f'由若依综合漏洞检测工具自动生成 · {gen_time} · 仅用于授权范围内的安全测试'
    )
    _set_run_font(footer_run, size=8, color=RGBColor(0x99, 0x99, 0x99))
    footer_para.paragraph_format.space_before = Pt(20)

    # 确保输出目录存在
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(out_path)
    return out_path
